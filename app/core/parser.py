# Steam Guide Downloader
# Copyright (c) 2025-2026 AlexAgents
# Licensed under the MIT License. See LICENSE file in the project root.

"""Steam guide parser"""

import os
import re
import logging
import threading
from typing import Callable, Optional

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

from app.config import AppConfig
from app.translations import get_text
from app.utils import clean_filename
from app.core.network import create_session, URLValidator, ImageCache
from app.core.docx_builder import DocxBuilder
from app.core.image_saver import ImageSaver
from app.core.pdf_converter import (
    convert_docx_to_pdf, check_available_converters
)

logger = logging.getLogger(__name__)


class GuideDownloader:
    def __init__(self, config: AppConfig):
        self.config = config
        self.session = create_session(config)
        self.image_cache = ImageCache(
            max_count=150,
            max_bytes=config.image_cache_limit_mb * 1024 * 1024
        )
        self._cancelled = threading.Event()
        self._page_size_mb = 0.0
        self._page_size_confirmed = False

    def cancel(self):
        self._cancelled.set()

    @property
    def is_cancelled(self):
        return self._cancelled.is_set()

    def confirm_large_page(self):
        """Called from GUI thread to confirm large page download."""
        self._page_size_confirmed = True

    def download(self, url, save_dir, lang_code, log_func,
                 finish_func, convert_pdf=False,
                 progress_func=None, save_images=False,
                 page_size_callback=None):
        self._cancelled.clear()
        self._page_size_confirmed = False
        self.image_cache.clear()

        try:
            self._do_download(
                url, save_dir, lang_code,
                log_func, convert_pdf, progress_func,
                save_images, page_size_callback
            )
        except Exception as e:
            logger.error(f"Download error: {e}", exc_info=True)
            log_func(f"Error: {e}")
        finally:
            logger.debug(self.image_cache.stats)
            self.image_cache.clear()
            try:
                self.session.close()
            except Exception:
                pass
            finish_func()

    def _do_download(self, url, save_dir, lang_code,
                     log_func, convert_pdf, progress_func,
                     save_images, page_size_callback):
        T = lambda key, *a: get_text(lang_code, key, *a)

        log_func(T("log_start", url))

        if convert_pdf:
            converters = check_available_converters()
            if not any(converters.values()):
                log_func(f"  {T('err_pdf_no_support')}")
                log_func("  Continuing with DOCX only\u2026")
                convert_pdf = False

        if not os.path.exists(save_dir):
            try:
                os.makedirs(save_dir, exist_ok=True)
            except OSError as e:
                log_func(f"{T('err_creating_dir')} {e}")
                return

        if self.is_cancelled:
            log_func(T("log_cancelled"))
            return

        try:
            response = self.session.get(
                url, timeout=self.config.timeout
            )
            response.raise_for_status()
            response.encoding = 'utf-8'
        except requests.ConnectionError:
            log_func(T("err_net_connection"))
            return
        except requests.Timeout:
            log_func(T("err_net_timeout"))
            return
        except requests.HTTPError as e:
            log_func(T("err_access", e.response.status_code))
            return
        except requests.RequestException as e:
            log_func(f"{T('err_net')} {e}")
            return

        if self.is_cancelled:
            log_func(T("log_cancelled"))
            return

        # Check page size — warn but don't block
        html_text = response.text
        page_size_mb = len(html_text.encode('utf-8')) / (1024 * 1024)
        self._page_size_mb = page_size_mb

        if (page_size_mb > self.config.max_page_size_mb
                and page_size_callback):
            # Signal GUI to ask user; wait for confirmation
            page_size_callback(page_size_mb)
            # Wait up to 60s for user response
            for _ in range(600):
                if self._page_size_confirmed or self.is_cancelled:
                    break
                import time
                time.sleep(0.1)

            if self.is_cancelled:
                log_func(T("log_cancelled"))
                return
            
            if not self._page_size_confirmed:
                log_func(T("log_cancelled"))
                return

        soup = BeautifulSoup(html_text, 'html.parser')

        doc = Document()
        self._setup_styles(doc)

        guide_title = self._extract_title(soup)
        doc.add_heading(guide_title, 0)

        safe_title = clean_filename(guide_title)
        if not safe_title or len(safe_title) < 2:
            gid = URLValidator.extract_guide_id(url)
            safe_title = f"manual_{gid}" if gid else "manual_unknown"

        full_path = os.path.join(save_dir, f"{safe_title}.docx")
        log_func(T("log_file_target", full_path))

        if self.is_cancelled:
            log_func(T("log_cancelled"))
            return

        # Create image saver if option enabled
        image_saver = None
        if save_images:
            image_saver = ImageSaver(
                save_dir=save_dir,
                guide_title=safe_title,
                config=self.config,
                session=self.session
            )

        builder = DocxBuilder(
            doc, config=self.config, session=self.session,
            image_cache=self.image_cache, log_func=log_func,
            image_saver=image_saver
        )

        content_found = self._process_content(
            soup, doc, builder, lang_code,
            log_func, progress_func
        )

        if not content_found and not self.is_cancelled:
            log_func(T("err_content"))
            return

        # Save DOCX — even if cancelled (partial save)
        try:
            doc.save(full_path)
            if self.is_cancelled:
                log_func(T("log_cancelled_partial", full_path))
            else:
                log_func(T("log_success", full_path))
        except PermissionError:
            log_func(T("err_permission"))
            return
        except OSError as e:
            log_func(f"Error: {e}")
            return

        # If cancelled, don't do post-processing
        if self.is_cancelled:
            return

        # Save images after DOCX is saved
        if image_saver and image_saver.count > 0:
            saved_count = image_saver.save_all(
                log_func=log_func,
                cancel_check=lambda: self.is_cancelled
            )
            log_func(T("log_images_total", saved_count))
            log_func(T("log_images_saved", image_saver.images_dir))

        if convert_pdf and not self.is_cancelled:
            log_func(T("log_pdf_converting"))
            success, result = convert_docx_to_pdf(
                full_path, log_func,
                timeout=self.config.pdf_timeout
            )
            if success:
                log_func(T("log_pdf_success", result))
            else:
                log_func(f"  {T('err_pdf_failed')}")
                log_func(result)

    def _setup_styles(self, doc):
        try:
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
        except Exception as e:
            logger.warning(f"Style error: {e}")

    def _extract_title(self, soup):
        title_node = soup.find('div', class_='workshopItemTitle')
        if title_node:
            t = title_node.get_text(strip=True)
            if t:
                return t
        title_tag = soup.find('title')
        if title_tag:
            t = title_tag.get_text(strip=True)
            t = re.sub(r'\s*::\s*Steam Community.*$', '', t)
            if t:
                return t
        return "Steam_Guide"

    def _process_content(self, soup, doc, builder,
                         lang_code, log_func,
                         progress_func=None):
        T = lambda key, *a: get_text(lang_code, key, *a)

        sections = soup.find_all(
            'div', class_='subSection detailBox'
        )

        if sections:
            total = len(sections)
            log_func(T("log_sections_found", total))

            for i, section in enumerate(sections):
                if self.is_cancelled:
                    return True  # partial content exists

                if progress_func:
                    progress_func(i + 1, total)

                title_div = section.find(
                    'div', class_='subSectionTitle'
                )
                if title_div:
                    ch = title_div.get_text(" ", strip=True)
                    doc.add_heading(ch, 1)
                    short = (
                        ch[:40] + ("\u2026" if len(ch) > 40 else "")
                    )
                    log_func(T("log_processing", short))

                desc_div = section.find(
                    'div', class_='subSectionDesc'
                )
                if desc_div:
                    for child in desc_div.children:
                        if self.is_cancelled:
                            return True
                        builder.process_node(child)
                    builder.close_paragraph()
            return True

        content = (
            soup.find('div', id='guideContent')
            or soup.find('div', class_='guide subSections')
        )

        if content:
            log_func(T("log_processing", "main content"))
            for child in content.children:
                if self.is_cancelled:
                    return True
                builder.process_node(child)
            builder.close_paragraph()
            return True

        return False
"""Tests for DocxBuilder components"""

import pytest
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.core.docx_builder import StyleContext, DocxBuilder
from app.config import AppConfig


class TestStyleContext:
    """Style context for DOCX formatting"""

    def test_default_all_false(self):
        ctx = StyleContext()
        assert not ctx.bold
        assert not ctx.italic
        assert not ctx.underline
        assert not ctx.strike
        assert not ctx.spoiler
        assert not ctx.code

    def test_set_bold(self):
        ctx = StyleContext(bold=True)
        assert ctx.bold
        assert not ctx.italic

    def test_set_multiple(self):
        ctx = StyleContext(bold=True, italic=True, code=True)
        assert ctx.bold
        assert ctx.italic
        assert ctx.code
        assert not ctx.underline

    def test_copy_independent(self):
        original = StyleContext(bold=True)
        copy = original.copy()
        copy.italic = True
        assert original.bold
        assert not original.italic
        assert copy.bold
        assert copy.italic

    def test_copy_preserves_all(self):
        original = StyleContext(
            bold=True, italic=True, underline=True,
            strike=True, spoiler=True, code=True
        )
        copy = original.copy()
        assert copy.bold
        assert copy.italic
        assert copy.underline
        assert copy.strike
        assert copy.spoiler
        assert copy.code

    def test_copy_does_not_affect_original(self):
        original = StyleContext()
        copy = original.copy()
        copy.bold = True
        copy.italic = True
        copy.underline = True
        assert not original.bold
        assert not original.italic
        assert not original.underline


class TestDocxBuilderBasic:
    """DocxBuilder HTML to DOCX conversion"""

    def _make_builder(self):
        from docx import Document
        doc = Document()
        config = AppConfig()
        builder = DocxBuilder(doc, config=config)
        return doc, builder

    def test_process_plain_text(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        html = "Hello World"
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        assert len(doc.paragraphs) >= 1
        assert "Hello World" in doc.paragraphs[0].text

    def test_process_bold_text(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        html = "<b>Bold text</b>"
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        assert len(doc.paragraphs) >= 1
        runs = doc.paragraphs[0].runs
        assert len(runs) >= 1
        assert runs[0].bold is True
        assert "Bold" in runs[0].text

    def test_process_italic_text(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        html = "<i>Italic text</i>"
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        runs = doc.paragraphs[0].runs
        assert len(runs) >= 1
        assert runs[0].italic is True

    def test_process_heading(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        html = "<h1>Title</h1>"
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        # Should have a heading paragraph
        found = False
        for p in doc.paragraphs:
            if "Title" in p.text:
                found = True
                break
        assert found

    def test_process_code_block(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        html = "<code>print('hello')</code>"
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        runs = doc.paragraphs[0].runs
        assert len(runs) >= 1
        assert runs[0].font.name == 'Courier New'

    def test_process_br_creates_new_paragraph(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        html = "Line 1<br>Line 2"
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) >= 2

    def test_process_list(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        html = "<ul><li>Item 1</li><li>Item 2</li></ul>"
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) >= 2
        assert "Item 1" in texts[0]

    def test_empty_block_no_crash(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        html = "<div></div><p></p><div><br></div>"
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        # Should not crash

    def test_nested_formatting(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        html = "<b><i>Bold and italic</i></b>"
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        runs = doc.paragraphs[0].runs
        assert len(runs) >= 1
        assert runs[0].bold is True
        assert runs[0].italic is True

    def test_max_recursion_depth(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        # Create deeply nested HTML
        html = "<div>" * 60 + "Deep" + "</div>" * 60
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        # Should not crash or stack overflow

    def test_strikethrough(self):
        from bs4 import BeautifulSoup
        doc, builder = self._make_builder()
        html = "<s>Strikethrough</s>"
        soup = BeautifulSoup(html, 'html.parser')
        for child in soup.children:
            builder.process_node(child)
        builder.close_paragraph()
        runs = doc.paragraphs[0].runs
        assert len(runs) >= 1
        assert runs[0].font.strike is True